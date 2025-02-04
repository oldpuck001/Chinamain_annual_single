# generate_audit_report.py

import os
import json
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side

def generate_audit_report(request):
    
    folder_path = request.get('data', {}).get('project_folder', '')

    settings_path = os.path.join(folder_path, '项目数据', 'settings.json')

    with open(settings_path, 'r', encoding='utf-8') as f:
        settings_dict = json.load(f)

    # 提取所需信息
    audit_opinion = request.get('data', {}).get('audit_opinion', '')
    report_date = request.get('data', {}).get('report_date', '')
    report_number = request.get('data', {}).get('report_number', '')
    report_version_notes = request.get('data', {}).get('report_version_notes', '')
    period = settings_dict['被审计会计期间']
    deadline = settings_dict['被审计会计报表截止日']
    accounting_firm = settings_dict['会计师事务所名称']
    enterprise_name = settings_dict['企业名称']
    date_of_establishment = settings_dict['成立日期']
    approval_date = settings_dict['核准日期']
    unified_social_credit_code = settings_dict['统一社会信用代码']
    registered_capital = settings_dict['注册资本']
    legal_representative = settings_dict['法定代表人']
    registered_address = settings_dict['注册地址']
    business_scope = settings_dict['经营范围']

    audit_report_path = os.path.join(folder_path, '审计报告', f'{enterprise_name}{period}审计报告{report_version_notes}')

    # 确保保存路径存在
    if not os.path.exists(audit_report_path):
        os.makedirs(audit_report_path, exist_ok=True)

    if audit_opinion == 'standard':
        base_dir = os.path.dirname(__file__)
        text_template_path = os.path.join(base_dir, 'audit_report_standard.txt')
    else:
        pass

    audit_report_text_path = os.path.join(audit_report_path, '审计报告正文.docx')

    audit_report_sheet_path = os.path.join(audit_report_path, '财务报表.xlsx')
    
    audit_report_notes_path = os.path.join(audit_report_path, '报表附注.docx')

    # 格式化日期为 'yyyy年m月d日' 的格式
    formatted_report_date = report_date[:4] + '年' + report_date[5:7] + '月' + report_date[8:] + '日'

    # 定义要替换的旧内容和新内容
    replacement_data = {
        '***公司名称***': enterprise_name,
        '***报告文号***': report_number,
        '***截止日***': deadline,
        '***期间***': period,
        '***出具日期***': formatted_report_date
    }

    with open(text_template_path, 'r', encoding='utf-8') as file:
        lines = [line.rstrip('\n') for line in file.readlines()]     # 读取所有行并去除换行符

    # 用于存储修改后的内容
    modified_lines = []
    # 遍历每一行进行替换
    for line in lines:
        # 对每行中的特定内容进行替换
        for old_content, new_content in replacement_data.items():
            line = line.replace(old_content, new_content)
        # 将修改后的行添加到新列表中
        modified_lines.append(line)

    # 创建文档对象
    doc = Document()

    # 报告文号的正则表达式
    report_number_pattern = r'\S*字\S*号$'

    # 将修改后的内容打印出来（或保存到新文件中）
    for modified_line in modified_lines:
        paragraph = doc.add_paragraph()                             # 添加第一个段落并设置字体和对齐方式
        run = paragraph.add_run(modified_line)        
        run.font.name = 'SimSun'                                    # 设置字体
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')     # 设置中文字体
        if modified_line == '审 计 报 告' or modified_line == '':
            run.font.size = Pt(25)                                  # 设置字号
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER     # 设置居中对齐
        elif re.search(report_number_pattern, modified_line):
            run.font.size = Pt(12)                                  # 设置字号
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER     # 设置居中对齐
        else:
            run.font.size = Pt(12)                                  # 设置字号
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT       # 设置靠左对齐

    # 添加页眉
    header = doc.sections[0].header
    header_paragraph = header.paragraphs[0]
    header_paragraph.text = "XXX会计师事务所"
    header_run = header_paragraph.runs[0]
    header_run.font.name = 'SimSun'
    header_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    header_run.font.size = Pt(12)
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE  # 设置分散对齐

    # 添加页脚
    footer = doc.sections[0].footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.text = "地址：……          联系电话：***-********"
    footer_run = footer_paragraph.runs[0]
    footer_run.font.name = 'SimSun'
    footer_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    footer_run.font.size = Pt(12)
    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # 保存文档
    doc.save(audit_report_text_path)


    # 创建一个新的工作簿
    wb = Workbook()
    ws = wb.active
    ws.title = '资产负债表'

    # 设置标题
    ws.merge_cells('A1:F1')
    ws['A1'].value = '资产负债表'
    ws['A1'].font = Font(size=24, bold=True)
    ws['A1'].alignment = Alignment(horizontal='center', vertical='center')

    # 设置报表信息
    ws.merge_cells('A2:B2')
    ws['A2'].value = f'编制单位：{enterprise_name}'
    ws['A2'].font = Font(size=12)
    ws['A2'].alignment = Alignment(horizontal='left', vertical='center')

    ws.merge_cells('C2:D2')
    ws['C2'].value = f'          {deadline}'
    ws['C2'].font = Font(size=12)
    ws['C2'].alignment = Alignment(horizontal='left', vertical='center')

    ws['F2'].value = f'单位：元'
    ws['F2'].font = Font(size=12)
    ws['F2'].alignment = Alignment(horizontal='left', vertical='center')

    # 设置列标题
    ws['A3'].value = '项目'
    ws['B3'].value = '期末余额'
    ws['C3'].value = '上年末余额'
    ws['D3'].value = '项目'
    ws['E3'].value = '期末余额'
    ws['F3'].value = '上年末余额'
    for row in ws['A3:F3']:
        for cell in row:
            cell.font = Font(bold=True)
            cell.alignment = Alignment(horizontal='center')
    
    # 报表项目
    balance_project_list = [['流动资产：', '流动负债：'],\
                            ['货币资金', '短期借款'],\
                            ['结算备付金*', '向中央银行借款*'],\
                            ['拆出资金*', '拆入资金*'],\
                            ['交易性金融资产', '交易性金融负债'],\
                            ['△以公允价值计量且其变动计入当期损益的金融资产', '△以公允价值计量且其变动计入当期损益的金融负债'],\
                            ['衍生金融资产', '衍生金融负债'],\
                            ['应收票据', '应付票据'],\
                            ['应收账款', '应付账款'],\
                            ['应收款项融资', '预收款项'],\
                            ['预付款项', '合同负债'],\
                            ['应收保费*', '卖出回购金融资产款*'],\
                            ['应收分保账款*', '吸收存款及同业存放*'],\
                            ['应收分保合同准备金*', '代理买卖证券款*'],\
                            ['其他应收款', '代理承销证券款*'],\
                            ['买入返售金融资产*', '应付职工薪酬'],\
                            ['存货', '应交税费'],\
                            ['合同资产', '其他应付款'],\
                            ['持有待售资产', '应付手续费及佣金*'],\
                            ['一年内到期的非流动资产', '应付分保账款*'],\
                            ['其他流动资产', '持有待售负债'],\
                            ['流动资产合计', '一年内到期的非流动负债'],\
                            ['非流动资产：', '其他流动负债'],\
                            ['发放贷款和垫款*', '流动负债合计'],\
                            ['债权投资', '非流动负债：'],\
                            ['△可供出售金融资产', '保险合同准备金*'],\
                            ['其他债权投资', '长期借款'],\
                            ['△持有至到期投资', '应付债券'],\
                            ['长期应收款', '应付债券：优先股'],\
                            ['长期股权投资', '应付债券：永续债'],\
                            ['其他权益工具投资', '租赁负债'],\
                            ['其他非流动金融资产', '长期应付款'],\
                            ['投资性房地产', '预计负债'],\
                            ['固定资产', '递延收益'],\
                            ['在建工程', '递延所得税负债'],\
                            ['生产性生物资产', '其他非流动负债'],\
                            ['油气资产', '非流动负债合计'],\
                            ['使用权资产', '负债合计'],\
                            ['无形资产', '所有者权益（或股东权益）：'],\
                            ['开发支出', '实收资本（或股本）'],\
                            ['商誉', '#减：已归还投资'],\
                            ['长期待摊费用', '#实收资本（或股本）净额'],\
                            ['递延所得税资产', '其他权益工具'],\
                            ['其他非流动资产', '其他权益工具：优先股'],\
                            ['非流动资产合计', '其他权益工具：永续债'],\
                            ['', '资本公积'],\
                            ['', '减：库存股'],\
                            ['', '其他综合收益'],\
                            ['', '专项储备'],\
                            ['', '盈余公积'],\
                            ['', '一般风险准备*'],\
                            ['', '未分配利润'],\
                            ['', '*归属于母公司所有者权益（或股东权益）合计'],\
                            ['', '*少数股东权益'],\
                            ['', '所有者权益（或股东权益）合计'],\
                            ['资产总计', '负债和所有者权益（或股东权益）总计']]

    start_row = 4
    for p in balance_project_list:
        left = p[0]
        if left == '':
            pass
        else:
            if left == '应付债券：优先股' or left == '其他权益工具：优先股':
                ws.cell(row=start_row, column=1, value='  其中：优先股')
            elif left == '应付债券：永续债' or left == '其他权益工具：永续债':
                ws.cell(row=start_row, column=1, value='        永续债')
            else:
                ws.cell(row=start_row, column=1, value=left)
        right = p[1]
        if right == '':
            pass
        else:
            if right == '应付债券：优先股' or right == '其他权益工具：优先股':
                ws.cell(row=start_row, column=4, value='  其中：优先股')
            elif right == '应付债券：永续债' or right == '其他权益工具：永续债':
                ws.cell(row=start_row, column=4, value='        永续债')
            else:
                ws.cell(row=start_row, column=4, value=right)
        start_row += 1

    # 设置单元格格式为数值，保留两位小数，并显示千分位符
    for row in ws['B4:C59']:
        for cell in row:
            cell.number_format = '#,##0.00'
    for row in ws['E4:F59']:
        for cell in row:
            cell.number_format = '#,##0.00'

    ws['A25'].font = Font(bold=True)
    ws['B25'].font = Font(bold=True)
    ws['C25'].font = Font(bold=True)
    ws['A48'].font = Font(bold=True)
    ws['B48'].font = Font(bold=True)
    ws['C48'].font = Font(bold=True)
    ws['A59'].font = Font(bold=True)
    ws['B59'].font = Font(bold=True)
    ws['C59'].font = Font(bold=True)

    ws['D27'].font = Font(bold=True)
    ws['E27'].font = Font(bold=True)
    ws['F27'].font = Font(bold=True)
    ws['D40'].font = Font(bold=True)
    ws['E40'].font = Font(bold=True)
    ws['F40'].font = Font(bold=True)
    ws['D41'].font = Font(bold=True)
    ws['E41'].font = Font(bold=True)
    ws['F41'].font = Font(bold=True)
    ws['D58'].font = Font(bold=True)
    ws['E58'].font = Font(bold=True)
    ws['F58'].font = Font(bold=True)
    ws['D59'].font = Font(bold=True)
    ws['E59'].font = Font(bold=True)
    ws['F59'].font = Font(bold=True)

    # 设置A9、D9、D59三个单元格的自动换行
    ws['A9'].alignment = Alignment(wrap_text=True)
    ws['D9'].alignment = Alignment(wrap_text=True)
    ws['D56'].alignment = Alignment(wrap_text=True)
    ws['D59'].alignment = Alignment(wrap_text=True)

    # 设置列宽
    ws.column_dimensions['A'].width = 28
    ws.column_dimensions['B'].width = 18
    ws.column_dimensions['C'].width = 18
    ws.column_dimensions['D'].width = 28
    ws.column_dimensions['E'].width = 18
    ws.column_dimensions['F'].width = 18

    # 添加边框样式
    thin_border = Border(left=Side(style='thin'),
                         right=Side(style='thin'),
                         top=Side(style='thin'),
                         bottom=Side(style='thin'))

    for row in ws.iter_rows(min_row=3, min_col=1, max_row=59, max_col=6):
        for cell in row:
            cell.border = thin_border

    # 保存报表文件
    wb.save(audit_report_sheet_path)


    # 创建文档对象
    notes_doc = Document()

    # 添加表格一個4x3的表格
    table = notes_doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'                                                  # 设置表格样式（默認樣式）

    # 填充表格内容并设置对齐方式
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.text = ''                                                      # 先清空默认文本
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(f'单元格({i+1}, {j+1})')
    
            run.font.name = 'SimSun'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
            run.font.size = Pt(12)
            if i == 0:
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER    # 第一行居中对齐
            else:
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT      # 其他行靠右对齐

    # 设置表格边框：
    # 上/下边框：4（默认 0.5 pt）
    # 右/左边框：隐藏
    # 内部横/竖边框：2（默认 0.25 pt）

    tbl = table._element  # 获取表格的 XML 结构

    # 查找或创建 <w:tblPr>（表格属性）
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    # 查找或创建 <w:tblBorders>（表格边框）
    tblBorders = tblPr.find(qn("w:tblBorders"))
    if tblBorders is None:
        tblBorders = OxmlElement("w:tblBorders")
        tblPr.append(tblBorders)

    # 创建边框元素并设置属性
    border_settings = {
        "top": {"sz": "8", "val": "single"},        # 上边框
        "bottom": {"sz": "8", "val": "single"},     # 下边框
        "right": {"val": "none"},                   # 右边框隐藏
        "left": {"val": "none"},                    # 左边框隐藏
        "insideH": {"sz": "4", "val": "single"},    # 内部水平边框
        "insideV": {"sz": "4", "val": "single"}     # 内部垂直边框
    }

    for border_name, attrs in border_settings.items():
        border = tblBorders.find(qn(f"w:{border_name}"))
        if border is None:
            border = OxmlElement(f"w:{border_name}")
            tblBorders.append(border)

        for key, value in attrs.items():
            border.set(qn(f"w:{key}"), value)  # 设置属性

    # 保存文檔
    notes_doc.save(audit_report_notes_path)

    # 返回成功信息
    return ['generate_audit_report']